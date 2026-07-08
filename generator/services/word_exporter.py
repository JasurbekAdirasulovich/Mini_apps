import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any

class WordExporter:
    """
    Generates professional Microsoft Word (.docx) files for exam variants.
    Includes separate pages for each variant and a consolidated Answer Key on the final page.
    """

    @staticmethod
    def generate_docx(variants_data: Dict[str, Any]) -> io.BytesIO:
        """
        Creates a DOCX document in-memory containing exam variants.

        Args:
            variants_data (dict): The dictionary containing subject, topics, and list of variants.

        Returns:
            io.BytesIO: Binary stream of the DOCX file.
        """
        doc = Document()
        
        # Configure default page margins (1 inch on all sides)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Set base styles (Calibri or Arial, 11pt, dark text)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        subject = variants_data.get('subject', 'General Subject')
        difficulty = variants_data.get('difficulty', 'medium').upper()
        language = variants_data.get('language', 'uz')
        variants = variants_data.get('variants', [])

        # ----------------------------------------------------
        # 1. GENERATE EXAM PAPERS FOR EACH VARIANT
        # ----------------------------------------------------
        for idx, var in enumerate(variants):
            var_name = var.get('name', f"Variant {idx+1}")
            questions = var.get('questions', [])

            # Professional Header block
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_p.add_run(f"{subject.upper()}\n")
            title_run.font.size = Pt(16)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37) # Dark Slate Gray

            sub_p = doc.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_p.add_run(f"{var_name.upper()} | Qiyinlik darajasi: {difficulty}\n")
            sub_run.font.size = Pt(12)
            sub_run.font.bold = True
            sub_run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

            # Date and Student details line
            details_table = doc.add_table(rows=1, cols=2)
            details_table.autofit = False
            details_table.columns[0].width = Inches(4.5)
            details_table.columns[1].width = Inches(2.0)

            # Left cell: Student Name
            cell_left = details_table.cell(0, 0)
            p_left = cell_left.paragraphs[0]
            p_left.add_run("Talaba F.I.Sh: __________________________________").font.bold = True
            
            # Right cell: Group/Date
            cell_right = details_table.cell(0, 1)
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.add_run("Guruh: _________").font.bold = True

            doc.add_paragraph() # Spacing

            # Write Questions
            for q_idx, q in enumerate(questions):
                q_text = q.get('question', '')
                options = q.get('options', {})

                # Question text
                q_p = doc.add_paragraph()
                q_run = q_p.add_run(f"{q_idx + 1}. {q_text}")
                q_run.font.bold = True
                
                # Option choices
                for opt_key in sorted(options.keys()):
                    opt_val = options[opt_key]
                    opt_p = doc.add_paragraph()
                    opt_p.paragraph_format.left_indent = Inches(0.3)
                    opt_p.add_run(f"{opt_key}) {opt_val}")

                # Spacer between questions
                doc.add_paragraph()

            # Separate each variant on its own page
            # Prevent empty page break on the very last variant before the Answer Key page
            doc.add_page_break()

        # ----------------------------------------------------
        # 2. GENERATE ANSWER KEYS PAGE (Oxirgi sahifada)
        # ----------------------------------------------------
        key_title_p = doc.add_paragraph()
        key_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        key_title_run = key_title_p.add_run("JAVOBLAR KALITI / ANSWER KEYS\n")
        key_title_run.font.size = Pt(16)
        key_title_run.font.bold = True
        key_title_run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81) # Green color

        info_p = doc.add_paragraph()
        info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_p.add_run(f"Fan: {subject}\nYaratilgan sana: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n").font.italic = True

        # Render keys as a clean table
        # 1 column for question index, then 1 column per variant
        cols_count = len(variants) + 1
        rows_count = len(variants[0].get('questions', [])) + 1 if variants else 1
        
        keys_table = doc.add_table(rows=rows_count, cols=cols_count)
        keys_table.style = 'Light Shading Accent 1' # Built-in table style

        # Table Header
        hdr_cells = keys_table.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run("T/r").font.bold = True
        for v_idx, var in enumerate(variants):
            var_name = var.get('name', f"Var {v_idx+1}")
            hdr_cells[v_idx + 1].paragraphs[0].add_run(var_name).font.bold = True

        # Fill Table Answers
        for r_idx in range(1, rows_count):
            row_cells = keys_table.rows[r_idx].cells
            row_cells[0].paragraphs[0].add_run(str(r_idx))
            
            for v_idx, var in enumerate(variants):
                q_list = var.get('questions', [])
                if r_idx - 1 < len(q_list):
                    ans = q_list[r_idx - 1].get('answer', '-')
                    row_cells[v_idx + 1].paragraphs[0].add_run(ans)

        # Save to file stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
